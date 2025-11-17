#!/usr/bin/env python3
"""
Read Word documents (.docx) and extract text content.
"""
import sys
import os
from pathlib import Path

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠ python-docx not installed. Install with: pip install python-docx")


def read_docx(file_path: str) -> str:
    """
    Read a .docx file and extract all text content.
    
    Parameters
    ----------
    file_path : str
        Path to the .docx file
        
    Returns
    -------
    str
        Extracted text content
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is required. Install with: pip install python-docx")
    
    doc = Document(file_path)
    text_content = []
    
    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_content.append(para.text)
    
    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            text_content.append(" | ".join(row_text))
    
    return "\n".join(text_content)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python read_word_doc.py <path_to_docx_file>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    
    if not os.path.exists(file_path):
        print(f"❌ File not found: {file_path}")
        sys.exit(1)
    
    if not file_path.endswith('.docx'):
        print(f"⚠ Warning: File doesn't have .docx extension: {file_path}")
    
    try:
        content = read_docx(file_path)
        print("=" * 70)
        print(f"📄 Content from: {file_path}")
        print("=" * 70)
        print(content)
        print("=" * 70)
        print(f"✓ Extracted {len(content)} characters")
    except Exception as e:
        print(f"❌ Error reading document: {e}")
        sys.exit(1)

